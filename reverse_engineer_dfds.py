#!/usr/bin/env python3
"""
Reverse-Engineer Example DFDs
Reads all department .docx files from example_dfds/ and generates:
  - data/reference_dfds/{slug}.json   (canonical DFD JSON)
  - data/reference_dfds/{slug}.html   (pre-rendered HTML DFD)

No LLM calls — pure Python table parsing.
Zero variance guaranteed across re-runs.
"""

import json
import os
import re
import sys

from docx import Document

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from agent.dfd_html_renderer import DFDHTMLRenderer
from agent.dfd_validator import validate_dfd, format_validation_report

EXAMPLE_DIR = os.path.join(os.path.dirname(__file__), "example_dfds")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "data", "reference_dfds")

# Color palette for dispersal sinks
COLORS = [
    "#1565c0",  # blue
    "#e91e63",  # pink
    "#4caf50",  # green
    "#7b1fa2",  # purple
    "#ff6f00",  # amber
    "#00838f",  # teal
    "#c62828",  # red
    "#6a1b9a",  # deep purple
    "#2e7d32",  # dark green
    "#ef6c00",  # orange
]


def slugify(name: str) -> str:
    """Convert department name to a filesystem-safe slug."""
    s = name.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = s.strip("_")
    return s


def extract_dept_info(doc: Document) -> dict:
    """Extract department name and org info from Table 1."""
    info = {}
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows:
            key = row.cells[0].text.strip()
            val = row.cells[1].text.strip()
            info[key] = val
    return info


def extract_key_processes(doc: Document) -> list:
    """Extract Key Processes from Table 2."""
    processes = []
    if len(doc.tables) > 2:
        t = doc.tables[2]
        for ri in range(1, len(t.rows)):
            row = t.rows[ri]
            activity = row.cells[1].text.strip() if len(row.cells) > 1 else ""
            observations = row.cells[2].text.strip() if len(row.cells) > 2 else ""
            if activity:
                # Parse activity: first line is the name, rest is description
                lines = activity.split("\n")
                name = lines[0].strip()
                description = " ".join(l.strip() for l in lines[1:] if l.strip())
                processes.append({
                    "name": name,
                    "description": description,
                    "observations": observations,
                })
    return processes


def extract_data_mapping(doc: Document) -> list:
    """Extract Data Mapping from Table 3."""
    data_map = []
    if len(doc.tables) > 3:
        t = doc.tables[3]
        headers = [c.text.strip() for c in t.rows[0].cells]
        for ri in range(1, len(t.rows)):
            row = t.rows[ri]
            entry = {}
            for ci, h in enumerate(headers):
                if ci < len(row.cells):
                    entry[h] = row.cells[ci].text.strip()
            data_map.append(entry)
    return data_map


def extract_introduction(doc: Document) -> str:
    """Extract introduction paragraph text."""
    intro_lines = []
    in_intro = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Introduction" in p.text:
            in_intro = True
            continue
        if in_intro:
            if p.style.name.startswith("Heading"):
                break
            if p.text.strip():
                intro_lines.append(p.text.strip())
    return " ".join(intro_lines)


def build_dfd_json(
    dept_name: str,
    key_processes: list,
    data_mapping: list,
    intro: str,
) -> dict:
    """
    Build a DFD JSON object from extracted docx data.
    
    Strategy:
    - Data categories → collection sources under business processes
    - Storage locations → storage_systems
    - Key processes → business processes under actors
    - Other departments mentioned → dispersal sinks
    """
    # ── Determine central process ──
    # Clean up multi-line dept names
    clean_dept = dept_name.split("\n")[0].strip()
    central_process = f"ISFC {clean_dept}"

    # ── Extract unique storage systems ──
    storage_set = set()
    for entry in data_mapping:
        sl = entry.get("Storage Location", "")
        # Split compound storage locations
        for part in re.split(r"[/,]\s*", sl):
            part = part.strip().rstrip("\n").strip()
            if part and len(part) > 1 and part.lower() not in ("na", "n/a", ""):
                storage_set.add(part)

    storage_systems = []
    for s in sorted(storage_set):
        stype = "cloud"
        s_lower = s.lower()
        if any(k in s_lower for k in ["one drive", "onedrive", "aws", "salesforce", "snowflake"]):
            stype = "cloud"
        elif any(k in s_lower for k in ["local", "physical", "branch", "endpoint", "end point"]):
            stype = "local"
        elif any(k in s_lower for k in ["database", "db", "finnone", "hris", "portal"]):
            stype = "database"
        storage_systems.append({"name": s, "type": stype})

    # ── Build business processes from key_processes + data_mapping ──
    # Group data categories by owner/purpose into business processes
    bp_list = []
    for i, proc in enumerate(key_processes):
        bp_id = f"bp_{slugify(proc['name'])[:20]}_{i}"
        # Find data items related to this process
        related_data = []
        for entry in data_mapping:
            purpose = entry.get("Purpose", "").lower()
            category = entry.get("Data Category", "")
            description = entry.get("Description", "")
            owner = entry.get("Data Owner", "").lower()
            proc_lower = proc["name"].lower()
            # Match data items to processes by keyword overlap
            if (any(word in purpose for word in proc_lower.split() if len(word) > 3) or
                any(word in owner for word in proc_lower.split() if len(word) > 3) or
                any(word in category.lower() for word in proc_lower.split() if len(word) > 3)):
                related_data.append({
                    "category": category,
                    "description": description,
                })

        # If no data matched, assign all data items to the first process
        if not related_data and i == 0:
            related_data = [
                {"category": e.get("Data Category", ""), "description": e.get("Description", "")}
                for e in data_mapping
            ]

        # Extract data elements from description
        data_elements = []
        for rd in related_data:
            if rd["category"]:
                data_elements.append(rd["category"])

        # If still empty, use process description keywords
        if not data_elements:
            data_elements = ["General Data"]

        collection_sources = [{
            "name": proc["name"],
            "data_elements": list(set(data_elements))[:8],
        }]

        bp_list.append({
            "id": bp_id,
            "name": proc["name"],
            "collection_sources": collection_sources,
        })

    # ── Build actors ──
    # External: customers (first 1-2 business processes are usually customer-facing)
    # Internal: department processes
    # Vendor: technology partners

    customer_bps = []
    internal_bps = []

    # Heuristic: customer-facing processes mention "customer", "call", "loan", "complaint"
    customer_keywords = {"customer", "call", "communication", "complaint", "grievance",
                         "loan sourcing", "channel", "feedback", "query", "ticket"}
    
    for bp in bp_list:
        bp_lower = bp["name"].lower()
        if any(kw in bp_lower for kw in customer_keywords):
            customer_bps.append(bp)
        else:
            internal_bps.append(bp)

    # Ensure at least one customer BP and one internal BP
    if not customer_bps and bp_list:
        customer_bps = [bp_list[0]]
        internal_bps = bp_list[1:]
    if not internal_bps and len(bp_list) > 1:
        internal_bps = [bp_list[-1]]

    actors = [
        {
            "id": "customers",
            "name": "Customers",
            "type": "external",
            "color": "#fffde7",
            "business_processes": customer_bps,
        },
        {
            "id": "internal",
            "name": "Internal Departments",
            "type": "internal",
            "color": "#fce4ec",
            "business_processes": internal_bps if internal_bps else [{
                "id": "bp_dept_ops",
                "name": f"{clean_dept} Operations",
                "collection_sources": [{
                    "name": "Department Data",
                    "data_elements": [e.get("Data Category", "") for e in data_mapping[:5]]
                }],
            }],
        },
        {
            "id": "vendors",
            "name": "Vendors/Partners",
            "type": "vendor",
            "color": "#f1f8e9",
            "business_processes": [],
        },
    ]

    # ── Build dispersal sinks ──
    # Common ISFC departments that appear as data receivers
    sink_keywords = {
        "credit": ("Credit Department", "internal"),
        "sales": ("Sales Department", "internal"),
        "operations": ("Operations Department", "internal"),
        "compliance": ("Compliance Department", "internal"),
        "audit": ("Internal Audit Department", "internal"),
        "field audit": ("Field Audit & Anti-Fraud", "internal"),
        "grievance": ("Grievance Department", "internal"),
        "hr": ("HR Department", "internal"),
        "legal": ("Legal Department", "internal"),
        "finance": ("Finance Department", "internal"),
        "it": ("IT Department", "internal"),
        "collection": ("Collections Department", "internal"),
        "marketing": ("Marketing Department", "internal"),
        "treasury": ("Treasury Department", "internal"),
        "secretarial": ("Secretarial Department", "internal"),
        "rbi": ("RBI / Regulatory Bodies", "customers"),
        "nhb": ("NHB", "customers"),
        "vendor": ("Third-Party Vendors", "vendors"),
        "partner": ("Technology Partners", "vendors"),
        "kalyera": ("Kalyera Exotel", "vendors"),
        "exotel": ("Kalyera Exotel", "vendors"),
    }

    dispersal_sinks = []
    seen_sinks = set()
    color_idx = 0
    
    # Scan introduction, observations, and data mapping for mentions of other depts
    all_text = intro.lower()
    for proc in key_processes:
        all_text += " " + proc.get("description", "").lower()
        all_text += " " + proc.get("observations", "").lower()
    for dm in data_mapping:
        all_text += " " + dm.get("Purpose", "").lower()
        all_text += " " + dm.get("Data Owner", "").lower()

    for kw, (sink_name, actor_id) in sink_keywords.items():
        # Don't add the department itself as a sink
        if kw in clean_dept.lower():
            continue
        if kw in all_text and sink_name not in seen_sinks:
            seen_sinks.add(sink_name)
            dispersal_sinks.append({
                "id": f"sink_{slugify(sink_name)[:20]}",
                "name": sink_name,
                "actor_id": actor_id,
                "color": COLORS[color_idx % len(COLORS)],
            })
            color_idx += 1

    # Ensure at least 2 sinks
    if len(dispersal_sinks) < 2:
        # Add generic sinks based on storage
        for sys in storage_systems[:3]:
            sname = f"Data to {sys['name']}"
            if sname not in seen_sinks:
                seen_sinks.add(sname)
                dispersal_sinks.append({
                    "id": f"sink_{slugify(sname)[:20]}",
                    "name": sname,
                    "actor_id": "internal",
                    "color": COLORS[color_idx % len(COLORS)],
                })
                color_idx += 1
                if len(dispersal_sinks) >= 3:
                    break

    # ── Build data flows ──
    data_flows = []
    flow_color_idx = 0

    # Inbound: each customer BP → central_process
    for bp in customer_bps:
        label = bp["collection_sources"][0]["data_elements"][0] if bp["collection_sources"] else ""
        data_flows.append({
            "from_id": bp["id"],
            "to_id": "central_process",
            "color": COLORS[flow_color_idx % len(COLORS)],
            "label": label[:25] if label else bp["name"][:20],
        })
        flow_color_idx += 1

    # Inbound: internal BPs → central_process
    for bp in internal_bps:
        if bp.get("collection_sources"):
            label = bp["collection_sources"][0]["data_elements"][0] if bp["collection_sources"][0].get("data_elements") else bp["name"]
        else:
            label = bp["name"]
        data_flows.append({
            "from_id": bp["id"],
            "to_id": "central_process",
            "color": COLORS[flow_color_idx % len(COLORS)],
            "label": label[:25],
        })
        flow_color_idx += 1

    # Outbound: central_process → each dispersal sink
    for sink in dispersal_sinks:
        data_flows.append({
            "from_id": "central_process",
            "to_id": sink["id"],
            "color": sink["color"],
            "label": "",
        })

    # ── Build citations ──
    citations = []

    # Cite business processes from Key Processes table
    for bp in bp_list:
        proc_match = next((p for p in key_processes if p["name"] == bp["name"]), None)
        if proc_match:
            citations.append({
                "element_id": bp["id"],
                "element_name": bp["name"],
                "source_section": "Key Processes (Table 2)",
                "source_text": proc_match["description"][:200] if proc_match["description"] else proc_match["name"],
                "source_type": "docx_table",
            })

    # Cite data elements from Data Mapping table
    for entry in data_mapping:
        cat = entry.get("Data Category", "")
        desc = entry.get("Description", "")
        purpose = entry.get("Purpose", "")
        if cat:
            citations.append({
                "element_id": f"data_{slugify(cat)[:20]}",
                "element_name": cat,
                "source_section": "Data Mapping (Table 3)",
                "source_text": f"{cat}: {desc[:100]}. Purpose: {purpose[:100]}",
                "source_type": "docx_table",
            })

    # Cite storage systems
    for sys_item in storage_systems[:6]:
        # Find which data category uses this storage
        for entry in data_mapping:
            if sys_item["name"].lower() in entry.get("Storage Location", "").lower():
                citations.append({
                    "element_id": f"storage_{slugify(sys_item['name'])[:20]}",
                    "element_name": sys_item["name"],
                    "source_section": "Data Mapping (Table 3) — Storage Location",
                    "source_text": f"{entry.get('Data Category', '')} is stored in {entry.get('Storage Location', '')}",
                    "source_type": "docx_table",
                })
                break

    # Cite dispersal sinks
    for sink in dispersal_sinks:
        citations.append({
            "element_id": sink["id"],
            "element_name": sink["name"],
            "source_section": "Key Processes — Observations / Data Mapping — Purpose",
            "source_text": f"Data shared with {sink['name']} as identified in process observations and data flow analysis",
            "source_type": "inferred",
        })

    return {
        "department": clean_dept,
        "version": "1.0",
        "central_process": central_process,
        "actors": actors,
        "dispersal_sinks": dispersal_sinks,
        "storage_systems": storage_systems[:6],
        "data_flows": data_flows,
        "citations": citations,
    }


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    renderer = DFDHTMLRenderer()

    skipped = []
    processed = []

    for fname in sorted(os.listdir(EXAMPLE_DIR)):
        if not fname.endswith(".docx"):
            continue
        # Skip Master and Risk reports
        if "Master" in fname or "Risk" in fname:
            skipped.append(fname)
            continue

        filepath = os.path.join(EXAMPLE_DIR, fname)
        print(f"\n{'='*60}")
        print(f"Processing: {fname}")
        print(f"{'='*60}")

        try:
            doc = Document(filepath)
        except Exception as e:
            print(f"  ERROR reading docx: {e}")
            skipped.append(fname)
            continue

        # Extract data
        dept_info = extract_dept_info(doc)
        dept_name = dept_info.get("Department Name", fname.replace(".docx", ""))
        key_processes = extract_key_processes(doc)
        data_mapping = extract_data_mapping(doc)
        intro = extract_introduction(doc)

        print(f"  Department: {dept_name}")
        print(f"  Key Processes: {len(key_processes)}")
        print(f"  Data Items: {len(data_mapping)}")

        # Build DFD JSON
        dfd_json = build_dfd_json(dept_name, key_processes, data_mapping, intro)

        # Validate
        validation = validate_dfd(dfd_json)
        report = format_validation_report(validation)
        print(f"  Validation: {report}")

        # Save JSON
        slug = slugify(dept_name.split("\n")[0])
        json_path = os.path.join(OUTPUT_DIR, f"{slug}.json")
        with open(json_path, "w") as f:
            json.dump(dfd_json, f, indent=2)
        print(f"  JSON: {json_path}")

        # Render HTML
        html_content = renderer.render(dfd_json)
        html_path = os.path.join(OUTPUT_DIR, f"{slug}.html")
        with open(html_path, "w") as f:
            f.write(html_content)
        print(f"  HTML: {html_path}")

        processed.append({
            "file": fname,
            "dept": dept_name.split("\n")[0],
            "slug": slug,
            "score": validation["score"],
            "passed": validation["passed"],
        })

    # Summary
    print(f"\n{'='*60}")
    print(f"SUMMARY")
    print(f"{'='*60}")
    print(f"Processed: {len(processed)}")
    print(f"Skipped:   {len(skipped)} ({', '.join(skipped)})")
    print()
    for p in processed:
        status = "✓ PASS" if p["passed"] else "✗ FAIL"
        print(f"  {p['dept']:<45} {status} ({p['score']}/100)  → {p['slug']}")
    print(f"\nOutput directory: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
