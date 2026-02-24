#!/usr/bin/env python3
"""
Generate Sample Transcripts for All Departments
Reads department .docx reports from example_dfds/ and generates
realistic interview transcripts deterministically (no LLM needed).

Each transcript simulates a privacy assessment interview covering:
- Department overview and hierarchy
- Key business processes
- Data types collected and stored
- Systems used
- Data sharing / dispersal

Output: data/sample_transcripts/{dept_slug}_transcript.txt
"""

import os
import re
import sys
import json
import random
from datetime import datetime, timedelta

from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

EXAMPLE_DIR = os.path.join(os.path.dirname(__file__), "example_dfds")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "data", "sample_transcripts")

# Interviewer names
INTERVIEWERS = ["Rajesh Kumar (K&S Digiprotect)", "Priya Sharma (K&S Digiprotect)"]
# Common SPOC fallback
SPOC_FALLBACK = "Department SPOC"

random.seed(42)  # Deterministic


def slugify(name: str) -> str:
    s = name.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    return s.strip("_")


def extract_docx_data(filepath: str) -> dict:
    """Extract all relevant data from a department docx."""
    doc = Document(filepath)

    # Department info from Table 1
    dept_name = ""
    org_name = ""
    assessment_date = ""
    location = ""
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows:
            key = row.cells[0].text.strip()
            val = row.cells[1].text.strip()
            if key == "Department Name":
                dept_name = val
            elif key == "Organization Name":
                org_name = val
            elif "Date" in key:
                assessment_date = val
            elif key == "Location":
                location = val

    # SPOC info
    spoc_name = ""
    paragraphs = doc.paragraphs
    for idx, p in enumerate(paragraphs):
        if "Department SPOC" in p.text:
            # Look for name in nearby paragraphs
            for offset in range(1, 4):
                if idx + offset < len(paragraphs):
                    next_text = paragraphs[idx + offset].text.strip()
                    if next_text and not next_text.startswith(("Contact", "Email", "Phone")):
                        spoc_name = next_text.replace("Department SPOC:", "").strip()
                        break
            break

    # Introduction text
    intro_lines = []
    in_intro = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Introduction" in p.text:
            in_intro = True
            continue
        if in_intro:
            if p.style.name.startswith("Heading"):
                break
            if p.text.strip():
                intro_lines.append(p.text.strip())
    intro = " ".join(intro_lines)

    # Key Processes from Table 2
    processes = []
    if len(doc.tables) > 2:
        t = doc.tables[2]
        for ri in range(1, len(t.rows)):
            row = t.rows[ri]
            activity = row.cells[1].text.strip() if len(row.cells) > 1 else ""
            observations = row.cells[2].text.strip() if len(row.cells) > 2 else ""
            remarks = row.cells[3].text.strip() if len(row.cells) > 3 else ""
            if activity:
                lines = activity.split("\n")
                name = lines[0].strip()
                desc = " ".join(l.strip() for l in lines[1:] if l.strip())
                processes.append({
                    "name": name,
                    "description": desc,
                    "observations": observations,
                    "remarks": remarks,
                })

    # Data Mapping from Table 3
    data_map = []
    if len(doc.tables) > 3:
        t = doc.tables[3]
        headers = [c.text.strip() for c in t.rows[0].cells]
        for ri in range(1, len(t.rows)):
            row = t.rows[ri]
            entry = {}
            for ci, h in enumerate(headers):
                if ci < len(row.cells):
                    entry[h] = row.cells[ci].text.strip()
            data_map.append(entry)

    # Recommendations
    recommendations = []
    in_rec = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Recommendation" in p.text:
            in_rec = True
            continue
        if in_rec:
            if p.style.name.startswith("Heading"):
                break
            if p.text.strip() and p.style.name == "List Paragraph":
                recommendations.append(p.text.strip())

    return {
        "dept_name": dept_name.split("\n")[0].strip(),
        "org_name": org_name,
        "assessment_date": assessment_date,
        "location": location,
        "spoc_name": spoc_name or SPOC_FALLBACK,
        "intro": intro,
        "processes": processes,
        "data_map": data_map,
        "recommendations": recommendations[:6],
    }


def generate_transcript(data: dict) -> str:
    """Generate a realistic interview transcript from extracted docx data."""
    dept = data["dept_name"]
    org = data["org_name"] or "India Shelter Finance Corporation Limited"
    spoc = data["spoc_name"]
    interviewer = random.choice(INTERVIEWERS)
    date = data["assessment_date"] or "April 2024"

    lines = []

    def add(speaker, text):
        lines.append(f"{speaker}: {text}")

    def blank():
        lines.append("")

    # ── Header ──
    lines.append(f"TRANSCRIPT — Privacy Impact Assessment Interview")
    lines.append(f"Organization: {org}")
    lines.append(f"Department: {dept}")
    lines.append(f"Date: {date}")
    lines.append(f"Interviewer: {interviewer}")
    lines.append(f"Interviewee: {spoc}")
    lines.append(f"{'='*70}")
    blank()

    # ── Introduction ──
    add(interviewer, f"Good morning. Thank you for taking the time for this interview. "
        f"We are conducting a privacy assessment of the {dept} as part of "
        f"ISFC's DPDPA 2023 compliance initiative. "
        f"Could you start by giving us an overview of your department?")
    blank()

    if data["intro"]:
        # Split intro into natural chunks
        intro_sentences = data["intro"].split(". ")
        chunk1 = ". ".join(intro_sentences[:3]) + "."
        chunk2 = ". ".join(intro_sentences[3:6]) + "." if len(intro_sentences) > 3 else ""
        add(spoc, chunk1)
        if chunk2:
            blank()
            add(spoc, chunk2)
    else:
        add(spoc, f"Sure. The {dept} is responsible for managing various operations "
            f"related to our core business functions at {org}.")
    blank()

    # ── Key Processes ──
    add(interviewer, f"Let's walk through the key processes in your department. "
        f"Can you describe each major activity and how data flows through it?")
    blank()

    for i, proc in enumerate(data["processes"]):
        add(spoc, f"{'Our first' if i == 0 else 'Another'} key process is "
            f"**{proc['name']}**. {proc['description'][:300]}")
        blank()

        if proc.get("observations") and proc["observations"] != "NA":
            add(interviewer, f"What are the key observations around {proc['name']}?")
            blank()
            add(spoc, f"Regarding {proc['name']}, {proc['observations'][:250]}")
            blank()

        if proc.get("remarks") and proc["remarks"].strip():
            add(spoc, f"I should also note: {proc['remarks'][:200]}")
            blank()

    # ── Data Collection & Storage ──
    add(interviewer, f"Now let's talk about the types of personal data your department collects. "
        f"What categories of data do you handle, and where is it stored?")
    blank()

    for entry in data["data_map"]:
        category = entry.get("Data Category", "")
        description = entry.get("Description", "")
        purpose = entry.get("Purpose", "")
        storage = entry.get("Storage Location", "")
        classification = entry.get("Data Classification*", entry.get("Data Classification", ""))
        retention = entry.get("Retention Period", "")
        legal_basis = entry.get("Legal Obligation", entry.get("Legal Basis", ""))

        if category:
            add(spoc, f"We collect **{category}**. {description[:200] if description else ''}")
            blank()
            if purpose:
                add(interviewer, f"What is the purpose of collecting {category}?")
                blank()
                add(spoc, f"The purpose is: {purpose[:200]}")
                blank()
            if storage:
                add(interviewer, f"Where is {category} stored?")
                blank()
                add(spoc, f"{category} is stored in {storage}. "
                    f"{'Its classification is ' + classification + '.' if classification else ''} "
                    f"{'The retention period is ' + retention + '.' if retention and retention != 'Not Defined' else 'Retention period is currently not formally defined.'} "
                    f"{'Legal basis: ' + legal_basis + '.' if legal_basis else ''}")
                blank()

    # ── Data Sharing ──
    add(interviewer, f"Does the {dept} share data with other departments or external parties?")
    blank()

    # Infer sharing from observations and processes
    sharing_mentions = set()
    all_text = " ".join(p.get("description", "") + " " + p.get("observations", "") for p in data["processes"])
    dept_keywords = {
        "sales": "Sales Department",
        "credit": "Credit Department",
        "operations": "Operations Department",
        "compliance": "Compliance Department",
        "audit": "Internal Audit",
        "grievance": "Grievance Department",
        "hr": "HR Department",
        "legal": "Legal Department",
        "finance": "Finance Department",
        "it": "IT Department",
        "collection": "Collections Department",
        "marketing": "Marketing Department",
        "customer care": "Customer Care Department",
        "rbi": "RBI",
        "nhb": "NHB",
    }
    for kw, dname in dept_keywords.items():
        if kw in all_text.lower() and kw not in dept.lower():
            sharing_mentions.add(dname)

    if sharing_mentions:
        add(spoc, f"Yes, we regularly share relevant data with: "
            f"{', '.join(list(sharing_mentions)[:5])}. "
            f"This is necessary for cross-departmental coordination and regulatory compliance.")
    else:
        add(spoc, f"We share data on a need-to-know basis with other departments "
            f"as required by our internal policies and regulatory requirements.")
    blank()

    # ── Systems ──
    storages = set()
    for entry in data["data_map"]:
        sl = entry.get("Storage Location", "")
        for part in re.split(r"[/,]\s*", sl):
            part = part.strip().rstrip("\n").strip()
            if part and len(part) > 1 and part.lower() not in ("na", "n/a", ""):
                storages.add(part)

    if storages:
        add(interviewer, f"What are the primary IT systems and tools used by the {dept}?")
        blank()
        add(spoc, f"We primarily use: {', '.join(sorted(storages))}. "
            f"These systems are managed by our IT department and "
            f"access is controlled through role-based access controls.")
        blank()

    # ── Consent & Compliance ──
    add(interviewer, f"How does the {dept} handle consent management for personal data processing?")
    blank()
    add(spoc, f"Consent is obtained as part of our standard onboarding and data collection processes. "
        f"However, I acknowledge that we need to formalize our consent management "
        f"mechanism in line with DPDPA 2023 requirements. This is something we're "
        f"actively working on with the Compliance team.")
    blank()

    # ── Recommendations discussion ──
    if data["recommendations"]:
        add(interviewer, f"Based on our assessment so far, we have some preliminary recommendations. "
            f"Let me share them with you.")
        blank()
        for rec in data["recommendations"][:4]:
            add(interviewer, f"Recommendation: {rec[:200]}")
            blank()
        add(spoc, f"Thank you. These are very helpful. We will take these into consideration "
            f"and work with our management to implement the necessary changes.")
        blank()

    # ── Closing ──
    add(interviewer, f"Thank you for your time. This has been very informative. "
        f"We'll compile our findings and share the assessment report with you.")
    blank()
    add(spoc, f"Thank you. Please let us know if you need any additional information.")
    blank()
    lines.append(f"{'='*70}")
    lines.append(f"[END OF TRANSCRIPT]")

    return "\n".join(lines)


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    results = []
    for fname in sorted(os.listdir(EXAMPLE_DIR)):
        if not fname.endswith(".docx"):
            continue
        if "Master" in fname or "Risk" in fname:
            continue

        filepath = os.path.join(EXAMPLE_DIR, fname)
        print(f"Processing: {fname}")

        try:
            data = extract_docx_data(filepath)
        except Exception as e:
            print(f"  ERROR: {e}")
            continue

        dept_name = data["dept_name"]
        slug = slugify(dept_name)

        transcript = generate_transcript(data)

        # Save transcript
        out_path = os.path.join(OUTPUT_DIR, f"{slug}_transcript.txt")
        with open(out_path, "w") as f:
            f.write(transcript)

        word_count = len(transcript.split())
        line_count = len(transcript.split("\n"))
        print(f"  → {slug}_transcript.txt  ({word_count} words, {line_count} lines)")

        results.append({
            "dept": dept_name,
            "slug": slug,
            "file": out_path,
            "words": word_count,
            "lines": line_count,
        })

    # Summary
    print(f"\n{'='*60}")
    print(f"SUMMARY — {len(results)} transcripts generated")
    print(f"{'='*60}")
    total_words = 0
    for r in results:
        print(f"  {r['dept']:<45} {r['words']:>5} words  {r['lines']:>4} lines")
        total_words += r["words"]
    print(f"\n  Total: {total_words} words across {len(results)} transcripts")
    print(f"  Output: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
